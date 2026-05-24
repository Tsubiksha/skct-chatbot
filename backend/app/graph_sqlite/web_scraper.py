"""
web_scraper.py — SKCT website crawler for Graph RAG.

Same-domain crawler for https://skct.edu.in/ using requests + BeautifulSoup.
Stores cleaned pages in SQLite, chunks them, indexes in FTS5.
Respects rate limits, depth, and page count.
"""

import re
import time
import logging
import zipfile
import json
import xml.etree.ElementTree as ET
from io import BytesIO
from collections import deque
from urllib.parse import urljoin, urlparse, urldefrag
from pathlib import Path

import requests
from bs4 import BeautifulSoup

from backend.app.config import settings
from backend.app.graph_sqlite.db import get_conn, table_count
from backend.app.graph_sqlite.text_cleaner import clean_html, detect_page_type, content_hash, is_useful
from backend.app.graph_sqlite.chunker import chunk_and_store

logger = logging.getLogger(__name__)

# ---- Skip patterns ----
_DOCUMENT_EXTENSIONS = {
    ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx",
}

_SKIP_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg", ".bmp", ".ico",
    ".css", ".js", ".json",
    ".mp4", ".mp3", ".wav", ".avi", ".mov",
    ".zip", ".rar", ".7z", ".tar", ".gz",
}

_SKIP_URL_PATTERNS = [
    r"/wp-json/", r"/wp-admin/", r"/xmlrpc", r"/feed/?$",
    r"/tag/", r"/category/", r"\?s=", r"/search",
    r"/media/\d+/?$",
    r"\?post_type=", r"10\.\d{4,9}/", r"\s", r"\(", r"\)",
    r"/login", r"/admin", r"/wp-login",
    r"google\.com/forms", r"docs\.google", r"forms\.gle",
    r"facebook\.com", r"twitter\.com", r"instagram\.com",
    r"youtube\.com", r"linkedin\.com", r"whatsapp",
    r"#", r"javascript:",
]

_SKIP_COMPILED = [re.compile(p, re.IGNORECASE) for p in _SKIP_URL_PATTERNS]

# Seed URLs to ensure key SKCT pages are visited
_SEED_URLS = [
    "https://skct.edu.in/",
    "https://skct.edu.in/principal/",
    "https://skct.edu.in/management/",
    "https://skct.edu.in/departments/",
    "https://skct.edu.in/placement/",
    "https://skct.edu.in/training/",
    "https://skct.edu.in/academics/",
    "https://skct.edu.in/events/",
    "https://skct.edu.in/research/",
    "https://skct.edu.in/facilities/",
    "https://skct.edu.in/industry-connect/",
    "https://skct.edu.in/admissions/",
    "https://skct.edu.in/examinations/",
]

_TRUSTED_EXTERNAL_SEEDS = [
    "https://www.tnea.in/CollegeInfo/2722.PDF",
]

_SESSION_HEADERS = {
    "User-Agent": "CollegeGraphRAGBot/1.0 educational local project",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate",
    "Connection": "keep-alive",
    "Cache-Control": "no-cache",
}

_SITEMAP_CANDIDATES = (
    "/sitemap.xml",
    "/sitemap_index.xml",
    "/page-sitemap.xml",
    "/post-sitemap.xml",
    "/wp-sitemap.xml",
)

_WORDPRESS_REST_TYPES = ("pages", "posts")


def _normalize_url(url: str, base: str = "") -> str | None:
    """Normalize URL: resolve relative, remove fragment, strip trailing slash."""
    try:
        if base:
            url = urljoin(base, url)
        url, _ = urldefrag(url)
        url = url.rstrip("/")
        if not url:
            return None
        return url
    except Exception:
        return None


def _is_same_domain(url: str, base_domain: str) -> bool:
    try:
        host = urlparse(url).netloc.lower().lstrip("www.")
        base = base_domain.lower().lstrip("www.")
        return host == base or host.endswith(f".{base}")
    except Exception:
        return False


def _should_skip(url: str) -> bool:
    """True if URL should be skipped."""
    parsed = urlparse(url)
    ext = Path(parsed.path).suffix.lower()
    if ext in _SKIP_EXTENSIONS:
        return True
    for pat in _SKIP_COMPILED:
        if pat.search(url):
            return True
    return False


def _is_document_url(url: str) -> bool:
    return Path(urlparse(url).path).suffix.lower() in _DOCUMENT_EXTENSIONS


def _candidate_from_srcset(value: str) -> list[str]:
    """Return URL candidates from an HTML srcset value."""
    candidates = []
    for part in value.split(","):
        raw = part.strip().split(" ", 1)[0].strip()
        if raw:
            candidates.append(raw)
    return candidates


def _extract_url_candidates(tag) -> list[str]:
    candidates = []
    for attr in ("href", "src", "data-src", "data-href", "data-url", "content"):
        raw = tag.get(attr)
        if raw:
            candidates.append(str(raw).strip())
    for attr in ("srcset", "data-srcset"):
        raw = tag.get(attr)
        if raw:
            candidates.extend(_candidate_from_srcset(str(raw)))
    return candidates


def _extract_links(soup: BeautifulSoup, base_url: str, base_domain: str) -> list[dict]:
    """Extract all same-domain internal links and same-domain downloadable documents."""
    links = []
    for tag in soup.find_all(True):
        for href in _extract_url_candidates(tag):
            if not href or href.startswith(("javascript:", "mailto:", "tel:")):
                continue
            normalized = _normalize_url(href, base_url)
            if not normalized:
                continue
            if not _is_same_domain(normalized, base_domain):
                continue
            if _should_skip(normalized):
                continue
            links.append({
                "url": normalized,
                "text": re.sub(r"\s+", " ", tag.get_text(" ", strip=True) or "").strip()[:250],
            })
    deduped = {}
    for link in links:
        deduped.setdefault(link["url"], link)
    return list(deduped.values())


def _root_url(start_url: str) -> str:
    parsed = urlparse(start_url)
    return f"{parsed.scheme}://{parsed.netloc}"


def _parse_sitemap_xml(content: bytes, sitemap_url: str) -> tuple[list[str], list[str]]:
    page_urls: list[str] = []
    nested_sitemaps: list[str] = []
    try:
        root = ET.fromstring(content)
    except Exception as e:
        logger.debug(f"[Scraper] Could not parse sitemap {sitemap_url}: {e}")
        return page_urls, nested_sitemaps

    for node in root.iter():
        if not node.tag.endswith("loc") or not node.text:
            continue
        loc = node.text.strip()
        if not loc:
            continue
        if loc.lower().endswith(".xml"):
            nested_sitemaps.append(loc)
        else:
            page_urls.append(loc)
    return page_urls, nested_sitemaps


def _discover_sitemap_urls(session: requests.Session, start_url: str,
                           base_domain: str, max_sitemaps: int = 50) -> list[str]:
    """Discover crawlable URLs from sitemap indexes and individual sitemaps."""
    root = _root_url(start_url)
    queue = deque(_normalize_url(path, root) for path in _SITEMAP_CANDIDATES)
    seen: set[str] = set()
    discovered: list[str] = []

    while queue and len(seen) < max_sitemaps:
        sitemap_url = queue.popleft()
        if not sitemap_url or sitemap_url in seen:
            continue
        seen.add(sitemap_url)
        if not _is_same_domain(sitemap_url, base_domain):
            continue
        try:
            resp = session.get(sitemap_url, timeout=15)
            if resp.status_code != 200:
                continue
            page_urls, nested = _parse_sitemap_xml(resp.content, sitemap_url)
            for loc in nested:
                n = _normalize_url(loc)
                if n and n not in seen:
                    queue.append(n)
            for loc in page_urls:
                n = _normalize_url(loc)
                if n and _is_same_domain(n, base_domain) and not _should_skip(n):
                    discovered.append(n)
        except Exception as e:
            logger.debug(f"[Scraper] Sitemap fetch failed {sitemap_url}: {e}")

    return list(dict.fromkeys(discovered))


def _discover_wordpress_urls(session: requests.Session, start_url: str,
                             base_domain: str, per_page: int = 100,
                             max_pages_per_type: int = 20) -> list[str]:
    """Discover public WordPress page/post links exposed by the REST API."""
    root = _root_url(start_url)
    discovered: list[str] = []

    for rest_type in _WORDPRESS_REST_TYPES:
        for page in range(1, max_pages_per_type + 1):
            api_url = f"{root}/wp-json/wp/v2/{rest_type}?per_page={per_page}&page={page}"
            try:
                resp = session.get(api_url, timeout=15)
                if resp.status_code != 200:
                    break
                items = resp.json()
                if not isinstance(items, list) or not items:
                    break
                for item in items:
                    if not isinstance(item, dict):
                        continue
                    guid = item.get("guid") if isinstance(item.get("guid"), dict) else {}
                    link = item.get("link") or guid.get("rendered")
                    n = _normalize_url(str(link or ""))
                    if n and _is_same_domain(n, base_domain) and not _should_skip(n):
                        discovered.append(n)
            except json.JSONDecodeError:
                break
            except Exception as e:
                logger.debug(f"[Scraper] WordPress REST fetch failed {api_url}: {e}")
                break

    return list(dict.fromkeys(discovered))


def _is_allowed_url(url: str, base_domain: str, trusted_external: set[str]) -> bool:
    return _is_same_domain(url, base_domain) or url in trusted_external


def _store_page_links(conn, page_id: int, source_url: str, links: list[dict], depth: int):
    for link in links:
        target = link.get("url")
        if not target or target == source_url:
            continue
        conn.execute(
            """INSERT INTO page_links (source_page_id, source_url, target_url, link_text, depth)
               VALUES (?, ?, ?, ?, ?)
               ON CONFLICT(source_url, target_url) DO UPDATE SET
                 source_page_id=excluded.source_page_id,
                 link_text=excluded.link_text,
                 depth=excluded.depth""",
            (page_id, source_url, target, link.get("text", ""), depth),
        )


def _extract_document_text(content: bytes, url: str, content_type: str = "") -> tuple[str, str]:
    """Extract text from downloaded PDF/Office files without writing them to disk."""
    ext = Path(urlparse(url).path).suffix.lower()
    title = Path(urlparse(url).path).name or url

    if ext == ".pdf" or "pdf" in content_type.lower():
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            pages = []
            for idx in range(len(doc)):
                text = doc[idx].get_text("text")
                if text and text.strip():
                    pages.append(f"Page {idx + 1}\n{text.strip()}")
            doc.close()
            return title, "\n\n".join(pages)
        except Exception as e:
            logger.debug(f"[Scraper] PyMuPDF failed for {url}: {e}")
            try:
                from pypdf import PdfReader
                reader = PdfReader(BytesIO(content))
                pages = []
                for idx, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(f"Page {idx + 1}\n{text.strip()}")
                return title, "\n\n".join(pages)
            except Exception as pdf_error:
                logger.warning(f"[Scraper] PDF extraction failed for {url}: {pdf_error}")
                return title, ""

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(BytesIO(content))
            parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return title, "\n".join(parts)
        except Exception as e:
            logger.warning(f"[Scraper] DOCX extraction failed for {url}: {e}")
            return title, ""

    if ext in {".xlsx", ".xls"}:
        try:
            import pandas as pd
            sheets = pd.read_excel(BytesIO(content), sheet_name=None, header=None)
            parts = []
            for sheet_name, df in sheets.items():
                df = df.dropna(how="all").dropna(axis=1, how="all")
                if not df.empty:
                    parts.append(f"Sheet: {sheet_name}")
                    parts.append(df.astype(str).to_csv(index=False, header=False))
            return title, "\n".join(parts)
        except Exception as e:
            logger.warning(f"[Scraper] spreadsheet extraction failed for {url}: {e}")
            return title, ""

    if ext == ".pptx":
        try:
            parts = []
            with zipfile.ZipFile(BytesIO(content)) as archive:
                slide_names = sorted(
                    name for name in archive.namelist()
                    if name.startswith("ppt/slides/slide") and name.endswith(".xml")
                )
                for idx, name in enumerate(slide_names, 1):
                    root = ET.fromstring(archive.read(name))
                    texts = [
                        node.text.strip()
                        for node in root.iter()
                        if node.tag.endswith("}t") and node.text and node.text.strip()
                    ]
                    if texts:
                        parts.append(f"Slide {idx}\n" + "\n".join(texts))
            return title, "\n\n".join(parts)
        except Exception as e:
            logger.warning(f"[Scraper] PPTX extraction failed for {url}: {e}")
            return title, ""

    return title, ""


def _is_document_response(url: str, content_type: str) -> bool:
    ct = content_type.lower()
    return _is_document_url(url) or any(token in ct for token in [
        "application/pdf",
        "wordprocessingml.document",
        "msword",
        "spreadsheetml.sheet",
        "vnd.ms-excel",
        "presentationml.presentation",
        "vnd.ms-powerpoint",
    ])


def _upsert_page(conn, url: str, title: str, page_type: str,
                 content: str, c_hash: str) -> tuple[int, str]:
    """
    Upsert page. Returns (page_id, status) where status is 'new'|'updated'|'unchanged'.
    """
    title = title or url
    page_type = page_type or "general"
    content = content or ""
    c_hash = c_hash or content_hash(content)

    existing = conn.execute(
        "SELECT id, title, page_type, content_hash FROM scraped_pages WHERE url=?", (url,)
    ).fetchone()

    if existing:
        if existing["content_hash"] == c_hash:
            if (existing["title"] or "") != (title or "") or (existing["page_type"] or "") != (page_type or ""):
                conn.execute(
                    """UPDATE scraped_pages
                       SET title=?, page_type=?, scraped_at=CURRENT_TIMESTAMP
                       WHERE id=?""",
                    (title, page_type, existing["id"]),
                )
            return existing["id"], "unchanged"
        conn.execute(
            """UPDATE scraped_pages
               SET title=?, page_type=?, content=?, content_hash=?, scraped_at=CURRENT_TIMESTAMP
               WHERE id=?""",
            (title, page_type, content, c_hash, existing["id"]),
        )
        return existing["id"], "updated"
    else:
        cur = conn.execute(
            """INSERT INTO scraped_pages
                 (url, title, page_type, content, content_hash, scraped_at)
               VALUES (?,?,?,?,?,CURRENT_TIMESTAMP)""",
            (url, title, page_type, content, c_hash),
        )
        return cur.lastrowid, "new"


def scrape_website(
    college_url: str | None = None,
    max_pages: int | None = None,
    max_depth: int | None = None,
    crawl_delay: float | None = None,
    force_reindex: bool = False,
) -> dict:
    """
    Crawl SKCT website, store cleaned pages, chunk and index in FTS5.

    Args:
        college_url:   Base URL (defaults to COLLEGE_URL env var)
        max_pages:     Max pages to store (defaults to GRAPH_RAG_MAX_PAGES)
        max_depth:     Max crawl depth (defaults to GRAPH_RAG_CRAWL_DEPTH)
        crawl_delay:   Seconds between requests (defaults to GRAPH_RAG_CRAWL_DELAY)
        force_reindex: Re-chunk even unchanged pages

    Returns:
        Summary dict.
    """
    # ---- Config ----
    start_url = (college_url or settings.COLLEGE_URL or "").strip().rstrip("/")
    if not start_url:
        return {
            "status": "error",
            "error": (
                "COLLEGE_URL is not configured. "
                "Please add COLLEGE_URL=https://skct.edu.in/ to your .env file."
            ),
        }

    max_pages   = max_pages if max_pages is not None else settings.GRAPH_RAG_MAX_PAGES
    max_depth   = max_depth if max_depth is not None else settings.GRAPH_RAG_CRAWL_DEPTH
    crawl_delay = crawl_delay if crawl_delay is not None else settings.GRAPH_RAG_CRAWL_DELAY
    base_domain = urlparse(start_url).netloc.lstrip("www.")
    trusted_external = {
        _normalize_url(seed) or seed
        for seed in _TRUSTED_EXTERNAL_SEEDS
    }

    logger.info(
        f"[Scraper] Starting crawl: url={start_url}, max_pages={max_pages}, "
        f"max_depth={max_depth}, delay={crawl_delay}s"
    )

    # ---- State ----
    visited:  set[str] = set()
    queue:    deque     = deque()   # (url, depth)
    stats = {
        "pages_visited": 0,
        "pages_saved": 0,
        "pages_updated": 0,
        "pages_unchanged": 0,
        "documents_saved": 0,
        "chunks_created": 0,
        "pages_skipped_short": 0,
        "errors": [],
        "skipped_external": 0,
    }

    # Seed with start URL + known good SKCT URLs
    normalized_start = _normalize_url(start_url) or start_url
    queue.append((normalized_start, 0))
    for seed in _SEED_URLS:
        n = _normalize_url(seed)
        if n and _is_same_domain(n, base_domain):
            queue.append((n, 1))  # depth 1 — discovered pages
    for seed in trusted_external:
        queue.append((seed, 1))

    # ---- HTTP session ----
    session = requests.Session()
    session.headers.update(_SESSION_HEADERS)

    # Warm-up request for cookies
    try:
        session.get(start_url, timeout=10)
    except Exception:
        pass

    discovered_seed_count = 0
    for discovered_url in _discover_sitemap_urls(session, start_url, base_domain):
        queue.append((discovered_url, 1))
        discovered_seed_count += 1
    for discovered_url in _discover_wordpress_urls(session, start_url, base_domain):
        queue.append((discovered_url, 1))
        discovered_seed_count += 1

    logger.info(f"[Scraper] Added {discovered_seed_count} discovered sitemap/API URLs")

    # ---- Crawl loop ----
    saved_count = 0
    conn = get_conn()

    try:
        while queue and saved_count < max_pages:
            url, depth = queue.popleft()
            if url in visited:
                continue
            if _should_skip(url):
                continue
            if not _is_allowed_url(url, base_domain, trusted_external):
                stats["skipped_external"] += 1
                continue
            if depth > max_depth:
                continue

            visited.add(url)
            stats["pages_visited"] += 1

            try:
                logger.info(f"[Scraper] Fetching ({depth}): {url}")
                resp = session.get(url, timeout=15, allow_redirects=True)

                # Check content type
                ct = resp.headers.get("Content-Type", "")
                is_html = "text/html" in ct or "application/xhtml" in ct
                is_doc = _is_document_response(resp.url or url, ct)
                if not is_html and not is_doc:
                    logger.debug(f"[Scraper] Non-HTML content-type: {ct} @ {url}")
                    continue

                if resp.status_code == 200:
                    html = ""
                    if is_doc:
                        title, cleaned = _extract_document_text(resp.content, resp.url or url, ct)
                    else:
                        html = resp.text
                        title, cleaned = clean_html(html, url)

                    if not is_useful(cleaned):
                        logger.debug(f"[Scraper] Too short, skipping: {url}")
                        stats["pages_skipped_short"] += 1
                        continue

                    page_type = "document" if is_doc else detect_page_type(url, title, cleaned)
                    c_hash    = content_hash(cleaned)

                    page_id, status = _upsert_page(
                        conn, url, title, page_type, cleaned, c_hash
                    )
                    conn.commit()

                    if status == "unchanged" and not force_reindex:
                        stats["pages_unchanged"] += 1
                    else:
                        n_chunks = chunk_and_store(page_id, url, title, page_type, cleaned)
                        stats["chunks_created"] += n_chunks
                        if status == "new":
                            stats["pages_saved"] += 1
                            if is_doc:
                                stats["documents_saved"] += 1
                        elif status == "updated":
                            stats["pages_updated"] += 1
                            if is_doc:
                                stats["documents_saved"] += 1
                        elif status == "unchanged":
                            stats["pages_unchanged"] += 1

                    saved_count += 1

                    # Discover links if within depth limit
                    if is_html and depth < max_depth:
                        soup = BeautifulSoup(html, "html.parser")
                        links = _extract_links(soup, url, base_domain)
                        _store_page_links(conn, page_id, url, links, depth + 1)
                        conn.commit()
                        for link in links:
                            link_url = link.get("url")
                            if link_url and link_url not in visited:
                                queue.append((link_url, depth + 1))

                elif resp.status_code in (301, 302, 307, 308):
                    logger.debug(f"[Scraper] Redirect: {url}")
                else:
                    logger.warning(f"[Scraper] HTTP {resp.status_code}: {url}")
                    stats["errors"].append(f"HTTP {resp.status_code}: {url}")

            except requests.exceptions.ConnectionError as e:
                msg = f"Connection error: {url}"
                logger.warning(f"[Scraper] {msg}: {e}")
                stats["errors"].append(msg)
            except requests.exceptions.Timeout:
                msg = f"Timeout: {url}"
                logger.warning(f"[Scraper] {msg}")
                stats["errors"].append(msg)
            except Exception as e:
                msg = f"Error fetching {url}: {e}"
                logger.error(f"[Scraper] {msg}")
                stats["errors"].append(msg)

            time.sleep(crawl_delay)

        conn.commit()

        # Log to DB
        conn.execute(
            """INSERT INTO ingestion_logs
                 (source_type, source_name, status, message, rows_inserted, rows_skipped, created_at)
               VALUES (?,?,?,?,?,?,CURRENT_TIMESTAMP)""",
            (
                "website", start_url, "success",
                (
                    f"Visited: {stats['pages_visited']}. "
                    f"Saved: {stats['pages_saved']}. "
                    f"Updated: {stats['pages_updated']}. "
                    f"Unchanged: {stats['pages_unchanged']}. "
                    f"Documents: {stats['documents_saved']}. "
                    f"Chunks: {stats['chunks_created']}. "
                    f"Errors: {len(stats['errors'])}."
                ),
                stats["pages_saved"] + stats["pages_updated"],
                stats["pages_skipped_short"],
            ),
        )
        conn.commit()

    except Exception as e:
        conn.rollback()
        logger.exception(f"[Scraper] Fatal error: {e}")
        stats["errors"].append(f"Fatal: {e}")
    finally:
        conn.close()
        session.close()

    return {
        "status": "success" if not stats["errors"] else "partial",
        "college_url": start_url,
        "pages_visited":  stats["pages_visited"],
        "pages_saved":    stats["pages_saved"],
        "pages_updated":  stats["pages_updated"],
        "pages_unchanged": stats["pages_unchanged"],
        "documents_saved": stats["documents_saved"],
        "chunks_created": stats["chunks_created"],
        "total_scraped_pages": table_count("scraped_pages"),
        "total_chunks":        table_count("website_chunks"),
        "errors": stats["errors"][:10],
    }
